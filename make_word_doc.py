from docx import Document
from datetime import date
from web_scrapper import WiktionaryBot as bot

doc = Document()
doc.add_heading('Latin Exercises', 0)

date = date.today().strftime('%d-%m-%Y')

tables_data = bot.get_words_formated_data(['hominum', 'exercitus', 'rosa'])


def make_to_fill_table(table_data):
    table_to_fill = doc.add_table(rows=0, cols=3) 

    for idx, (case, singular, plural) in enumerate(table_data):
        row = table_to_fill.add_row().cells
        row[0].text = case
        if idx == 0:
            row[1].text = singular
            row[2].text = plural
        else:
            row[1].text = ''
            row[2].text = ''


def make_filled_table(table_data):
    table_filled = doc.add_table(rows=0, cols=3) 
    
    for case, singular, plural in table_data: 
        row = table_filled.add_row().cells 
        row[0].text = case 
        row[1].text = singular 
        row[2].text = plural


for data in tables_data:
    make_to_fill_table(data)
    doc.add_paragraph()

doc.add_paragraph()

for data in tables_data:
    make_filled_table(data)
    doc.add_paragraph()


doc.save(f'./word_docs/{date}.docx') 


