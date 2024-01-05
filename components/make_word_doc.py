from docx import Document
from datetime import date

doc = Document()


def make_to_fill_table(table_data):
    table_to_fill = doc.add_table(rows=0, cols=3) 

    for idx, (case, singular, plural) in enumerate(table_data):
        row = table_to_fill.add_row().cells
        row[0].text = case
        if idx == 0:
            row[1].text = singular
            row[2].text = plural
        else:
            row[1].text = ''
            row[2].text = ''


def make_filled_table(table_data):
    table_filled = doc.add_table(rows=0, cols=3) 
    
    for case, singular, plural in table_data: 
        row = table_filled.add_row().cells 
        row[0].text = case 
        row[1].text = singular 
        row[2].text = plural


def build_document(tables_data):
    doc.add_heading('Latin Exercises', 0)
    

    for data in tables_data:
        make_to_fill_table(data)
        doc.add_paragraph()

    doc.add_paragraph()

    for data in tables_data:
        make_filled_table(data)
        doc.add_paragraph()

    today_date = date.today().strftime('%d-%m-%Y')
    doc.save(f'./word_docs/{today_date}.docx') 


